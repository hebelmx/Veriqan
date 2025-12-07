"""Generate realistic DOCX + CNBV-standard PDF/PNG fixtures."""

from __future__ import annotations

import hashlib
import json
import random
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pdf2image import convert_from_path
from PIL import Image, ImageDraw, ImageEnhance, ImageFilter
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas as canvas_module
from reportlab.platypus import (
    Image as RLImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .authority_templates import AuthorityTemplate, match_template


class FixtureRenderer:
    """Render metadata into multi-format fixtures using authority templates."""

    def __init__(self, output_dir: Optional[Path], fmt: str = "png", seed: Optional[int] = None) -> None:
        self.output_dir = output_dir
        self.format = fmt.lower()
        self.random = random.Random(seed)
        if self.output_dir:
            self.output_dir.mkdir(parents=True, exist_ok=True)

    def render(self, record_id: str, metadata: Dict[str, object], sections: Dict[str, str]) -> Dict[str, str]:
        if not self.output_dir:
            return {}

        artifacts: Dict[str, str] = {}
        template = match_template(str(metadata.get("autoridadEmisora", "")))
        docx_path = self._build_docx(record_id, metadata, sections, template)
        artifacts["docx"] = str(docx_path)

        pdf_path = self._build_cnbv_pdf(record_id, metadata, sections)
        artifacts["pdf"] = str(pdf_path)

        if self.format in {"png", "both"}:
            png_path = self._pdf_to_png(pdf_path)
            if png_path:
                artifacts["png"] = str(png_path)

        xml_path = self.output_dir / f"{record_id}.xml"
        self._write_xml(xml_path, metadata, template)
        artifacts["xml"] = str(xml_path)
        return artifacts

    def _record_hash(self, metadata: Dict[str, object]) -> str:
        raw = json.dumps(metadata, ensure_ascii=False, sort_keys=True)
        return hashlib.sha256(raw.encode("utf-8")).hexdigest().upper()

    def _build_docx(
        self,
        record_id: str,
        metadata: Dict[str, object],
        sections: Dict[str, str],
        template: AuthorityTemplate,
    ) -> Path:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        for line in template.title_lines:
            p = doc.add_paragraph(line.upper())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(12)
            run.bold = True

        subtitle = doc.add_paragraph(template.subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.italic = True
        subtitle.runs[0].font.size = Pt(10)

        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = "Table Grid"
        info_pairs = [
            ("Autoridad", metadata.get("autoridadEmisora", "")),
            ("Expediente", metadata.get("expediente", "")),
            ("Fecha", metadata.get("fecha", "")),
        ]
        for idx, (label, value) in enumerate(info_pairs):
            info_table.cell(idx, 0).text = label
            info_table.cell(idx, 1).text = str(value)

        doc.add_paragraph("")
        doc.add_paragraph(f"Tipo de requerimiento: {metadata.get('tipoRequerimiento', '')}").bold = True

        doc.add_paragraph(sections.get("autoridad", ""), style="Normal")
        doc.add_paragraph(sections.get("instrucciones", ""))
        doc.add_paragraph(sections.get("apercibimiento", ""))

        detalle = metadata.get("detalle") or {}
        descripcion = detalle.get("descripcion")
        if descripcion:
            doc.add_paragraph("Detalle:").runs[0].bold = True
            doc.add_paragraph(str(descripcion))

        doc.add_paragraph(
            f"Plazo de cumplimiento: {metadata.get('plazoDescripcion', '')} "
            f"({metadata.get('plazoDias', '')} días)."
        )
        doc.add_paragraph(f"Fundamento legal: {metadata.get('fundamentoLegal', '')}")

        signature = doc.add_paragraph("\n\n")
        signature.add_run(template.seal_text).bold = True
        doc.add_paragraph("_______________________________")
        doc.add_paragraph("Secretaría de Acuerdos / Autoridad competente")

        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = template.footer
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        docx_path = self.output_dir / f"{record_id}.docx"
        doc.save(docx_path)
        return docx_path

    def _logo_path(self) -> Optional[Path]:
        guess = Path("Prisma/Fixtures/PRP1/LogoMexico.jpg")
        if guess.exists():
            return guess
        alt = Path("Prisma/Code/Src/CSharp/Python/LogoMexico.jpg")
        if alt.exists():
            return alt
        return None

    def _build_cnbv_pdf(self, record_id: str, metadata: Dict[str, object], sections: Dict[str, str]) -> Path:
        pdf_path = self.output_dir / f"{record_id}.pdf"
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CNBVTitle", parent=styles["Title"], alignment=1, fontSize=14, leading=16)
        subtitle_style = ParagraphStyle("CNBVSubtitle", parent=styles["Normal"], alignment=1, fontSize=10, italic=True)
        body_style = styles["Normal"]
        body_style.fontSize = 10
        body_style.leading = 14
        small_style = ParagraphStyle("Small", parent=body_style, fontSize=9, leading=12)

        solicitante = metadata.get("autoridadEmisora", "Solicitante no especificado")
        requerimiento_id = (
            metadata.get("numeroIdentificacion")
            or metadata.get("solicitudSiara")
            or metadata.get("expediente")
            or f"REQ-{record_id}"
        )

        data_table = [
            ["Número de oficio CNBV", metadata.get("expediente", "N/D")],
            ["Autoridad solicitante", solicitante],
            ["Fecha de recepción", metadata.get("fecha", "")],
            ["Tipo de requerimiento", metadata.get("tipoRequerimiento", "")],
        ]

        partes = metadata.get("partes") or []
        partes_rows = [[idx + 1, entry] for idx, entry in enumerate(partes)]
        if not partes_rows:
            partes_rows = [["-", "No especificado"]]
        detalle = metadata.get("detalle") or {}
        story: List[object] = []
        logo_path = self._logo_path()
        if logo_path:
            logos = [RLImage(str(logo_path), width=1.1 * inch, height=0.45 * inch) for _ in range(5)]
        else:
            logos = [Paragraph("Gobierno de México", small_style) for _ in range(5)]
        logos_table = Table([logos], colWidths=[1.1 * inch] * 5)
        story.append(logos_table)
        story.append(Spacer(1, 0.1 * inch))

        story.append(Paragraph(f"Solicitante: {solicitante}", body_style))
        story.append(Spacer(1, 0.05 * inch))

        id_table = Table(
            [
                ["", Paragraph("No. de identificación del requerimiento", small_style)],
                ["", Paragraph(str(requerimiento_id), body_style)],
            ],
            colWidths=[4.0 * inch, 2.0 * inch],
        )
        id_table.setStyle(
            TableStyle(
                [
                    ("BOX", (1, 0), (1, 0), 0.5, colors.black),
                    ("BOX", (1, 1), (1, 1), 0.5, colors.black),
                    ("ALIGN", (1, 0), (1, 1), "CENTER"),
                    ("VALIGN", (1, 0), (1, 1), "MIDDLE"),
                ]
            )
        )
        story.append(id_table)
        story.append(Spacer(1, 0.1 * inch))

        story.append(Paragraph("COMISIÓN NACIONAL BANCARIA Y DE VALORES", title_style))
        story.append(Paragraph("Dirección General de Supervisión de Autoridades", subtitle_style))
        story.append(Spacer(1, 0.2 * inch))

        info_table = Table(data_table, colWidths=[2.5 * inch, 3.5 * inch])
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f2f2")),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.black),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ]
            )
        )
        story.append(info_table)
        story.append(Spacer(1, 0.2 * inch))

        story.append(Paragraph("Datos Generales del solicitante", body_style))
        solicitante_block = (
            f"{solicitante}<br/>"
            f"{detalle.get('domicilio', 'Domicilio no proporcionado')}<br/>"
            f"Área: {metadata.get('areaSolicitante', 'Departamento no especificado')}"
        )
        contacto_block = (
            f"Tipo de requerimiento: {metadata.get('tipoRequerimiento', 'N/D')}<br/>"
            f"Contacto: {metadata.get('contactoCNBV', 'Oficial de enlace CNBV')}<br/>"
            f"Teléfono: {metadata.get('telefonoCNBV', '55 1234 5678')}<br/>"
            f"Correo: {metadata.get('correoCNBV', 'atencion.autoridades@cnbv.gob.mx')}"
        )
        general_table = Table(
            [[Paragraph(solicitante_block, body_style), Paragraph(contacto_block, body_style)]],
            colWidths=[3.6 * inch, 2.4 * inch],
        )
        general_table.setStyle(
            TableStyle(
                [
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.black),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(general_table)
        story.append(Spacer(1, 0.2 * inch))

        story.append(Paragraph("Resumen del oficio remitido:", body_style))
        story.append(Paragraph(sections.get("autoridad", ""), body_style))
        story.append(Spacer(1, 0.15 * inch))
        story.append(Paragraph("Instrucciones relevantes:", body_style))
        story.append(Paragraph(sections.get("instrucciones", ""), body_style))
        story.append(Spacer(1, 0.15 * inch))
        story.append(Paragraph("Apercibimientos:", body_style))
        story.append(Paragraph(sections.get("apercibimiento", ""), body_style))
        story.append(Spacer(1, 0.2 * inch))

        descripcion = detalle.get("descripcion") or "Sin descripción detallada."
        story.append(Paragraph("Detalle transcrito:", body_style))
        story.append(Paragraph(str(descripcion), body_style))
        story.append(Spacer(1, 0.2 * inch))

        partes_table = Table([["#", "Parte involucrada"]] + partes_rows, colWidths=[0.6 * inch, 4.4 * inch])
        partes_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e6f0ff")),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.black),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ]
            )
        )
        story.append(Paragraph("Partes identificadas:", body_style))
        story.append(partes_table)
        story.append(Spacer(1, 0.2 * inch))

        facultades_text = metadata.get("facultadesAutoridad") or detalle.get("facultades") or (
            "La autoridad solicitante acredita facultades conferidas por la normatividad aplicable para requerir "
            "información financiera, así como asegurar el seguimiento de medidas preventivas."
        )
        story.append(Paragraph("Facultades de la autoridad:", body_style))
        story.append(Paragraph(facultades_text, body_style))
        story.append(Spacer(1, 0.2 * inch))

        fundamento_text = metadata.get("fundamentoLegal") or "Sin fundamento especificado."
        story.append(Paragraph("Fundamentos del requerimiento:", body_style))
        story.append(Paragraph(fundamento_text, body_style))
        story.append(Spacer(1, 0.2 * inch))

        missing_fields = [key for key, value in metadata.items() if value in (None, "", [], {}) and key not in {"detalle"}]
        if missing_fields:
            story.append(Paragraph("Campos sin información proporcionada:", body_style))
            story.append(Paragraph(", ".join(sorted(missing_fields)), body_style))
            story.append(Spacer(1, 0.2 * inch))

        story.append(Paragraph("Observaciones CNBV:", body_style))
        story.append(
            Paragraph(
                "Este documento es transcripción fiel del oficio recibido. "
                "CNBV no interpreta ni modifica la información; únicamente la valida y redistribuye a las instituciones financieras.",
                body_style,
            )
        )

        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=LETTER,
            leftMargin=0.8 * inch,
            rightMargin=0.8 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
        )

        hash_label = self._record_hash(metadata)[:18]

        def _watermark(canvas: canvas_module.Canvas, _doc: SimpleDocTemplate) -> None:
            canvas.saveState()
            canvas.setFont("Helvetica-Bold", 48)
            canvas.setFillColor(colors.Color(1, 0, 0, alpha=0.12))
            canvas.rotate(30)
            for y in range(-2, 8):
                canvas.drawString(-1.5 * inch, y * inch, f"{hash_label} CNBV CONFIDENCIAL")
                canvas.drawString(4 * inch, y * inch, f"CNBV {hash_label}")
            canvas.restoreState()

        doc.build(story, onFirstPage=_watermark, onLaterPages=_watermark)
        return pdf_path

    def _apply_scan_artifacts(self, image: Image.Image) -> Image.Image:
        # additive noise
        if self.random.random() > 0.4:
            noise = Image.new("L", image.size, 0)
            draw = ImageDraw.Draw(noise)
            count = int(image.width * image.height * 0.05)
            for _ in range(count):
                x = self.random.randint(0, image.width - 1)
                y = self.random.randint(0, image.height - 1)
                draw.point((x, y), self.random.randint(0, 255))
            noise = noise.filter(ImageFilter.GaussianBlur(1.0))
            image = Image.blend(image, Image.merge("RGB", (noise, noise, noise)), 0.1)

        if self.random.random() > 0.5:
            image = image.filter(ImageFilter.GaussianBlur(radius=self.random.uniform(0.3, 1.0)))

        if self.random.random() > 0.5:
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(self.random.uniform(0.8, 1.2))

        if self.random.random() > 0.5:
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(self.random.uniform(0.9, 1.1))

        if self.random.random() > 0.4:
            image = image.rotate(self.random.uniform(-0.8, 0.8), expand=True, fillcolor="white")
        return image

    def _pdf_to_png(self, pdf_path: Path) -> Optional[Path]:
        png_path = pdf_path.with_suffix(".png")
        try:
            images = convert_from_path(str(pdf_path), dpi=200)
        except Exception:
            return None
        if not images:
            return None
        processed = self._apply_scan_artifacts(images[0].convert("RGB"))
        processed.save(png_path, "PNG")
        return png_path

    def _write_xml(self, path: Path, metadata: Dict[str, object], template: AuthorityTemplate) -> None:
        root = ET.Element("Requerimiento")
        root.set("template", template.identifier)

        ordered_fields = [
            "fecha",
            "autoridadEmisora",
            "expediente",
            "tipoRequerimiento",
            "subtipoRequerimiento",
            "fundamentoLegal",
            "motivacion",
            "partes",
            "detalle",
        ]

        for field in ordered_fields:
            value = metadata.get(field)
            if field == "partes":
                partes_node = ET.SubElement(root, "partes")
                if value:
                    for parte in value:
                        parte_node = ET.SubElement(partes_node, "parte")
                        parte_node.text = str(parte)
                else:
                    partes_node.text = ""
            elif field == "detalle":
                detalle_node = ET.SubElement(root, "detalle")
                detalle = value or {}
                for sub_field in ["descripcion", "monto", "moneda", "activoVirtual"]:
                    child = ET.SubElement(detalle_node, sub_field)
                    child_value = detalle.get(sub_field)
                    child.text = "" if child_value in (None, "") else str(child_value)
            else:
                child = ET.SubElement(root, field)
                child.text = "" if value in (None, "") else str(value)

        extras = {
            "plazoDias": metadata.get("plazoDias"),
            "plazoDescripcion": metadata.get("plazoDescripcion"),
            "aseguramiento": metadata.get("aseguramiento"),
            "profileId": metadata.get("profileId"),
            "promptHints": metadata.get("promptHints"),
        }
        extra_node = ET.SubElement(root, "metadataAdicional")
        for key, value in extras.items():
            child = ET.SubElement(extra_node, key)
            if isinstance(value, list):
                for item in value:
                    list_child = ET.SubElement(child, "item")
                    list_child.text = str(item)
                if not value:
                    child.text = ""
            else:
                child.text = "" if value in (None, "") else str(value)

        tree = ET.ElementTree(root)
        tree.write(path, encoding="utf-8", xml_declaration=True)
