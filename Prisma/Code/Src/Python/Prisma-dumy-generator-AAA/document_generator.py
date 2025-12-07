# document_generator.py

import xml.etree.ElementTree as ET
from faker import Faker
import random
import datetime
import os
import argparse
from tqdm import tqdm
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

# Import the client and the new orchestrator
from ollama_client import get_llm_narrative, construct_junior_lawyer_prompt
from ollama_orchestrator import ensure_ollama_ready

# Initialize Faker for Mexican Spanish
fake = Faker('es_MX')

def generate_document_data(ollama_api_url: str, llm_model: str) -> dict:
    """Generates a dictionary of fake data for a single document package."""
    
    company_name = fake.company()
    oficio_year = str(datetime.date.today().year)
    area_clave = random.choice([3, 5])
    area_desc = "ASEGURAMIENTO" if area_clave == 3 else "INFORMACION"

    data = {
        "Cnbv_NumeroOficio": f"222/{fake.random_element(elements=('AAA', 'BBB', 'CCC'))}/{fake.random_number(digits=10, fix_len=True)}/{oficio_year}",
        "Cnbv_NumeroExpediente": f"A/AS1-{fake.random_number(digits=4, fix_len=True)}-{fake.random_number(digits=6, fix_len=True)}-{fake.random_element(elements=('AAA', 'BBB', 'CCC'))}",
        "Cnbv_SolicitudSiara": f"AGAFADAFSON2/{oficio_year}/{fake.random_number(digits=6, fix_len=True)}",
        "Cnbv_Folio": str(fake.random_number(digits=4)),
        "Cnbv_OficioYear": oficio_year,
        "Cnbv_AreaClave": str(area_clave),
        "Cnbv_AreaDescripcion": area_desc,
        "Cnbv_FechaPublicacion": datetime.date.today().isoformat(),
        "Cnbv_DiasPlazo": str(random.choice([3, 5, 7, 10])),
        "AutoridadNombre": fake.company() + " " + fake.company_suffix(),
        "Referencia2": f"IMSSCOB/{fake.random_number(digits=2)}/{fake.random_number(digits=2)}/{fake.random_number(digits=6)}/{oficio_year}",
        "TieneAseguramiento": "true" if area_desc == "ASEGURAMIENTO" else "false",
        "Nombre": company_name,
        "NombreCompleto": f"{company_name}, S.A. DE C.V.",
        "Rfc": fake.rfc(natural=False),
        "Domicilio": f"{fake.street_address()} CP {fake.postcode()} Col. {fake.random_element(elements=('Centro', 'Roma Norte', 'Condesa', 'Polanco', 'Juárez'))}, {fake.city()}",
        "Complementarios": f"Y{fake.random_number(digits=2)} W {fake.random_number(digits=5)} {fake.random_number(digits=2)}"
    }
    
    # Generate narrative text
    prompt = construct_junior_lawyer_prompt({"company_name": company_name})
    data["InstruccionesCuentasPorConocer"] = get_llm_narrative(ollama_api_url, prompt, model=llm_model)
    
    return data

def create_xml_expediente(data: dict):
    """Generates the XML expediente from a data dictionary."""
    cnbv_namespace = "http://www.cnbv.gob.mx"
    def nstag(tag):
        return f"{{{cnbv_namespace}}}{tag}"
    ET.register_namespace('', cnbv_namespace)
    
    root = ET.Element(nstag("Expediente"), {"xmlns:xsi": "http://www.w3.org/2001/XMLSchema-instance", "xmlns:xsd": "http://www.w3.org/2001/XMLSchema"})
    
    for key, value in data.items():
        if key not in ["InstruccionesCuentasPorConocer", "NombreCompleto", "Rfc", "Domicilio", "Complementarios"]:
            ET.SubElement(root, nstag(key)).text = value

    solicitud_partes = ET.SubElement(root, nstag("SolicitudPartes"))
    ET.SubElement(solicitud_partes, nstag("ParteId")).text = "1"
    ET.SubElement(solicitud_partes, nstag("Caracter")).text = "Patrón Determinado"
    ET.SubElement(solicitud_partes, nstag("Persona")).text = "Moral"
    ET.SubElement(solicitud_partes, nstag("Paterno"))
    ET.SubElement(solicitud_partes, nstag("Materno"))
    ET.SubElement(solicitud_partes, nstag("Nombre")).text = data["Nombre"]
    ET.SubElement(solicitud_partes, nstag("Rfc")).text = " " * 13

    solicitud_especifica = ET.SubElement(root, nstag("SolicitudEspecifica"))
    ET.SubElement(solicitud_especifica, nstag("SolicitudEspecificaId")).text = "1"
    ET.SubElement(solicitud_especifica, nstag("InstruccionesCuentasPorConocer")).text = data["InstruccionesCuentasPorConocer"]

    personas_solicitud = ET.SubElement(solicitud_especifica, nstag("PersonasSolicitud"))
    ET.SubElement(personas_solicitud, nstag("PersonaId")).text = "1"
    ET.SubElement(personas_solicitud, nstag("Caracter")).text = "Patrón Determinado"
    ET.SubElement(personas_solicitud, nstag("Persona")).text = "Moral"
    ET.SubElement(personas_solicitud, nstag("Paterno"))
    ET.SubElement(personas_solicitud, nstag("Materno"))
    ET.SubElement(personas_solicitud, nstag("Nombre")).text = data["NombreCompleto"]
    ET.SubElement(personas_solicitud, nstag("Rfc")).text = data["Rfc"]
    ET.SubElement(personas_solicitud, nstag("Relacion"))
    ET.SubElement(personas_solicitud, nstag("Domicilio")).text = data["Domicilio"]
    ET.SubElement(personas_solicitud, nstag("Complementarios")).text = data["Complementarios"]

    return ET.ElementTree(root)

def create_docx_document(data: dict, output_path: str):
    """Creates a .docx file from the generated data."""
    doc = Document()
    doc.add_heading(f"Requerimiento: {data['Cnbv_NumeroOficio']}", level=1)
    doc.add_paragraph(f"Fecha: {data['Cnbv_FechaPublicacion']}")
    doc.add_paragraph(f"Autoridad: {data['AutoridadNombre']}")
    doc.add_heading("Partes Involucradas", level=2)
    doc.add_paragraph(f"Nombre: {data['NombreCompleto']}\nRFC: {data['Rfc']}\nDomicilio: {data['Domicilio']}")
    doc.add_heading("Instrucciones", level=2)
    doc.add_paragraph(data["InstruccionesCuentasPorConocer"])
    doc.save(output_path)

def create_pdf_document(data: dict, output_path: str):
    """Creates a .pdf file from the generated data."""
    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    story = []
    
    story.append(Paragraph(f"Requerimiento: {data['Cnbv_NumeroOficio']}", styles['h1']))
    story.append(Paragraph(f"Fecha: {data['Cnbv_FechaPublicacion']}", styles['Normal']))
    story.append(Paragraph(f"Autoridad: {data['AutoridadNombre']}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Partes Involucradas", styles['h2']))
    story.append(Paragraph(f"Nombre: {data['NombreCompleto']}", styles['Normal']))
    story.append(Paragraph(f"RFC: {data['Rfc']}", styles['Normal']))
    story.append(Paragraph(f"Domicilio: {data['Domicilio']}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Instrucciones", styles['h2']))
    # Use a different style for the body to handle newlines better
    body_style = styles['BodyText']
    body_style.wordWrap = 'CJK'
    story.append(Paragraph(data["InstruccionesCuentasPorConocer"].replace('\n', '<br/>'), body_style))
    
    doc.build(story)

def generate_documents(count: int, output_dir: str, ollama_api_url: str, llm_model: str):
    """Generates a batch of complete document packages."""
    os.makedirs(output_dir, exist_ok=True)
    print(f"Generating {count} document package(s) in '{output_dir}'...")

    for i in tqdm(range(count)):
        try:
            data = generate_document_data(ollama_api_url, llm_model)
            base_filename = f"expediente_{i+1:03d}"
            
            # Create XML
            xml_tree = create_xml_expediente(data)
            xml_path = os.path.join(output_dir, f"{base_filename}.xml")
            xml_tree.write(xml_path, encoding='utf-8', xml_declaration=True)
            
            # Create DOCX
            docx_path = os.path.join(output_dir, f"{base_filename}.docx")
            create_docx_document(data, docx_path)

            # Create PDF
            pdf_path = os.path.join(output_dir, f"{base_filename}.pdf")
            create_pdf_document(data, pdf_path)

        except Exception as e:
            print(f"\nError generating document package {i+1}: {e}")
            continue

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Autonomous Dummy Document Generator")
    parser.add_argument("--count", type=int, default=1, help="Number of documents to generate.")
    parser.add_argument("--output", type=str, default="generated_documents", help="Directory to save the generated documents.")
    parser.add_argument("--model", type=str, default="llama3", help="Name of the Ollama model to use.")
    args = parser.parse_args()

    print("--- Starting Autonomous Dummy Document Generation ---")
    
    try:
        api_url = ensure_ollama_ready(model=args.model)
        generate_documents(
            count=args.count,
            output_dir=args.output,
            ollama_api_url=api_url,
            llm_model=args.model
        )
        print(f"\nSUCCESS: Batch generation complete. Files are in '{args.output}'.")
    except Exception as e:
        print(f"\nFATAL ERROR: The generation pipeline failed. Reason: {e}", file=sys.stderr)

    print("\n--- Pipeline Finished ---")
