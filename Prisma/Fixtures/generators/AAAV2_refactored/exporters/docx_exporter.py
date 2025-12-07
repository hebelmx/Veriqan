"""DOCX exporter using python-docx library."""

import base64
import io
from pathlib import Path
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DOCXExporter:
    """Export CNBV documents as Microsoft Word (.docx) format."""

    def __init__(self, logo_path: Optional[Path] = None):
        """Initialize DOCX exporter.

        Args:
            logo_path: Path to logo image file
        """
        self.logo_path = logo_path

    def export(self, data: Dict, output_path: Path) -> Path:
        """Export data to DOCX file.

        Args:
            data: Dictionary with document data
            output_path: Path for output DOCX file

        Returns:
            Path to generated DOCX file
        """
        doc = Document()

        # Set document margins (narrower)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Add watermark (background text)
        self._add_watermark(doc, data.get('Cnbv_SolicitudSiara', 'DOCUMENTO'))

        # Add logo header (5 repetitions)
        if self.logo_path and self.logo_path.exists():
            self._add_logo_header(doc)

        # Add document header text
        self._add_header_text(doc, data)

        # Add ID box (as a right-aligned bordered paragraph)
        self._add_id_box(doc, data.get('Cnbv_SolicitudSiara', 'N/A'))

        # Add separator
        doc.add_paragraph('_' * 80)

        # Add recipient
        self._add_recipient(doc, data)

        # Add sections
        self._add_section_header(doc, "Datos generales del solicitante")
        self._add_two_column_table(doc, data)

        self._add_section_header(doc, "Facultades de la Autoridad")
        self._add_paragraph_text(doc, data.get('FacultadesTexto', ''))

        self._add_section_header(doc, "Fundamento del Requerimiento")
        self._add_paragraph_text(doc, data.get('FundamentoTexto', ''))

        self._add_section_header(doc, "Motivación del requerimiento")
        self._add_paragraph_text(doc, data.get('MotivacionTexto', ''))
        self._add_paragraph_text(doc, data.get('MontoTexto', ''))

        self._add_section_header(doc, "Origen del requerimiento")
        self._add_origen_section(doc, data)

        self._add_section_header(doc, "Solicitudes específicas: 1")
        self._add_subsection_header(doc, "Personas de quien se requiere información")
        self._add_personas_table(doc, data)

        self._add_section_header(doc, "Cuentas por conocer")
        self._add_paragraph_text(doc, data.get('SectoresBancarios', ''))

        self._add_section_header(doc, "Instrucciones sobre las cuentas por conocer")
        self._add_paragraph_text(doc, data.get('InstruccionesCuentasPorConocer', ''))

        # Add closing text
        closing = f"Derivado de lo anterior solicitó a la comisión nacional bancaria y de valores sea atendido el presente requerimiento y gestionado por medio de del sistema de atención de requerimientos autoridad (SIARA) contando con el folio {data.get('Cnbv_SolicitudSiara', 'N/A')}"
        self._add_paragraph_text(doc, closing)

        # Add signature
        self._add_signature(doc, data)

        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return output_path

    def _add_watermark(self, doc: Document, text: str) -> None:
        """Add watermark to document.

        Args:
            doc: Document object
            text: Watermark text
        """
        # Watermarks in python-docx require XML manipulation
        # This is a simplified version - full implementation would need header/footer manipulation
        # For now, we'll skip the diagonal watermark in DOCX (works better in PDF)
        pass

    def _add_logo_header(self, doc: Document) -> None:
        """Add logo header with 5 repetitions.

        Args:
            doc: Document object
        """
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add 5 logos
        for i in range(5):
            run = paragraph.add_run()
            run.add_picture(str(self.logo_path), width=Inches(0.6))
            if i < 4:  # Add space between logos
                run.add_text(' ')

    def _add_header_text(self, doc: Document, data: Dict) -> None:
        """Add header text lines.

        Args:
            doc: Document object
            data: Document data
        """
        lines = [
            "Administración General de Auditoría Fiscal Federal",
            data.get('UnidadSolicitante', ''),
        ]

        for line in lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.size = Pt(11)

    def _add_id_box(self, doc: Document, folio: str) -> None:
        """Add ID box (right-aligned, bordered).

        Args:
            doc: Document object
            folio: Document folio/ID
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = p.add_run(f"No. De Identificación del Requerimiento\n")
        run.font.size = Pt(9)

        run = p.add_run(folio)
        run.bold = True
        run.font.size = Pt(11)

        # Add border (simplified - full implementation would use paragraph borders)
        # For now, just make it bold and prominent

    def _add_recipient(self, doc: Document, data: Dict) -> None:
        """Add recipient section.

        Args:
            doc: Document object
            data: Document data
        """
        lines = [
            data.get('Destinatario_Nombre', ''),
            data.get('Destinatario_Cargo', ''),
            data.get('Destinatario_Institucion', ''),
            data.get('Destinatario_Direccion', ''),
            "P r e s e n t e",
        ]

        for line in lines:
            p = doc.add_paragraph(line)
            if "P r e s e n t e" in line:
                p.runs[0].bold = True

    def _add_section_header(self, doc: Document, title: str) -> None:
        """Add bordered section header.

        Args:
            doc: Document object
            title: Section title
        """
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(11)

        # Add background color (gray)
        # This requires XML manipulation in python-docx
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F5F5F5')
        p._p.get_or_add_pPr().append(shading_elm)

    def _add_subsection_header(self, doc: Document, title: str) -> None:
        """Add subsection header (lighter background).

        Args:
            doc: Document object
            title: Subsection title
        """
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(10)

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F9F9F9')
        p._p.get_or_add_pPr().append(shading_elm)

    def _add_paragraph_text(self, doc: Document, text: str) -> None:
        """Add regular paragraph text.

        Args:
            doc: Document object
            text: Paragraph text
        """
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.size = Pt(11)

    def _add_two_column_table(self, doc: Document, data: Dict) -> None:
        """Add two-column table with solicitor info.

        Args:
            doc: Document object
            data: Document data
        """
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        row = table.rows[0]

        # Left column
        left_cell = row.cells[0]
        left_text = f"""ADMINISTRACIÓN GENERAL DE AUDITORÍA FISCAL FEDERAL
{data.get('UnidadSolicitante', '')}

Mesa, Turno y/o Unidad, Secretaría etc: {data.get('UnidadSolicitante', '')}
{data.get('DomicilioSolicitante', '')}"""
        left_cell.text = left_text

        # Right column
        right_cell = row.cells[1]
        right_text = f"""Requerimiento Hacendario

{data.get('ServidorPublico_Nombre', '')}
{data.get('ServidorPublico_Cargo', '')}
Tel. {data.get('ServidorPublico_Telefono', '')}
Correo: {data.get('ServidorPublico_Correo', '')}"""
        right_cell.text = right_text

    def _add_origen_section(self, doc: Document, data: Dict) -> None:
        """Add origen del requerimiento section.

        Args:
            doc: Document object
            data: Document data
        """
        fields = [
            ("¿Esta solicitud contiene requerimientos de aseguramiento?", data.get('TieneAseguramiento', 'No')),
            ("No de oficio de requerimiento", data.get('NoOficioRevision', 'N/A')),
            ("Monto a crédito", data.get('MontoCredito', 'N/A')),
            ("Créditos fiscales", data.get('CreditosFiscales', 'N/A')),
            ("Periodo de revisión", data.get('Periodos', 'N/A')),
        ]

        for label, value in fields:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)

    def _add_personas_table(self, doc: Document, data: Dict) -> None:
        """Add personas table.

        Args:
            doc: Document object
            data: Document data
        """
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Nombre', 'RFC', 'Carácter', 'Dirección', 'Datos complementarios']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Data row
        data_cells = table.rows[1].cells
        data_values = [
            data.get('Persona_Nombre', ''),
            data.get('Persona_Rfc', ''),
            data.get('Persona_Caracter', ''),
            data.get('Persona_Domicilio', ''),
            data.get('Persona_Complementarios', ''),
        ]

        for i, value in enumerate(data_values):
            data_cells[i].text = value

    def _add_signature(self, doc: Document, data: Dict) -> None:
        """Add signature section.

        Args:
            doc: Document object
            data: Document data
        """
        doc.add_paragraph()  # Space
        doc.add_paragraph()

        p = doc.add_paragraph("_" * 40)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph(data.get('ServidorPublico_Nombre', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        p = doc.add_paragraph(data.get('ServidorPublico_Cargo', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
